# -*- coding: utf-8 -*-
"""
Generator BAB IV & BAB V Skripsi SIAKAD NUJA.
Meniru format/style/layout dari screenshots/PROPOSAL-SKRIPSI.docx,
namun seluruh isi ditulis khusus untuk sistem SIAKAD Nurul Jadid.
"""
import os
from docx import Document
from docx.shared import Emu, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHOT = "screenshots"
OUT = "BAB-IV-V-SIAKAD-NUJA.docx"

# ---- konstanta format hasil analisis template ----
PAGE_W = 7560945
PAGE_H = 10693400
M_L, M_R, M_T, M_B = 1440180, 1080135, 1080135, 1080135
FONT = "Times New Roman"
SZ_BODY = Emu(152400)      # 12pt
SZ_CAP = Emu(139700)       # 11pt
INDENT = Emu(171450)       # first line indent body
IMG_W = Emu(5029200)       # 5.5 inch
H3_FIRST = Emu(-457200)
H3_LEFT = Emu(628650)

doc = Document()

# ---- section ----
sec = doc.sections[0]
sec.page_width = Emu(PAGE_W); sec.page_height = Emu(PAGE_H)
sec.left_margin = Emu(M_L); sec.right_margin = Emu(M_R)
sec.top_margin = Emu(M_T); sec.bottom_margin = Emu(M_B)

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = SZ_BODY
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(0)


def _set_font(style, bold=None, size=None):
    style.font.name = FONT
    if size is not None:
        style.font.size = size
    if bold is not None:
        style.font.bold = bold
    style.paragraph_format.line_spacing = 1.5


for nm, bold, size, before, after in [
    ("Heading 1", True, SZ_BODY, Pt(0), None),
    ("Heading 2", True, SZ_BODY, Emu(228600), None),
    ("Heading 3", True, SZ_BODY, Emu(152400), None),
]:
    st = doc.styles[nm]
    _set_font(st, bold, size)
    st.paragraph_format.space_before = before
    if after is not None:
        st.paragraph_format.space_after = after

cap = doc.styles["Caption"]
_set_font(cap, None, SZ_CAP)
cap.paragraph_format.space_before = Emu(152400)
cap.paragraph_format.space_after = Emu(38100)


# ---- helper builders ----
def h1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    return p


def h2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    return p


def h3(text):
    p = doc.add_paragraph(style="Heading 3")
    p.paragraph_format.first_line_indent = H3_FIRST
    p.paragraph_format.left_indent = H3_LEFT
    p.add_run(text)
    return p


def body(text, indent=True):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = INDENT
    p.add_run(text)
    return p


def module(letter, title):
    """Paragraf pembuka modul: ' a.  Judul Modul' (tebal)."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = INDENT
    r = p.add_run("%s.  %s" % (letter, title))
    r.bold = True
    return p


def figure(fname, caption):
    """Sisipkan gambar (center) + caption."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    path = os.path.join(SHOT, fname)
    run.add_picture(path, width=IMG_W)
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run(caption)


def caption_only(text):
    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run(text)


def _cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = SZ_CAP
    r.font.bold = bold


def spec_table(rows):
    """Tabel spesifikasi 2 kolom (Normal Table)."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (a, b) in enumerate(rows):
        _cell_text(t.rows[i].cells[0], a, bold=(i == 0))
        _cell_text(t.rows[i].cells[1], b, bold=(i == 0))
    return t


def blackbox_table(rows):
    """Tabel pengujian black box 6 kolom (Table Grid)."""
    header = ["No", "Kasus Uji", "Masukan", "Keluaran yang Diharapkan",
              "Keluaran yang Dihasilkan", "Status"]
    t = doc.add_table(rows=len(rows) + 1, cols=6)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htext in enumerate(header):
        _cell_text(t.rows[0].cells[j], htext, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.CENTER if j in (0, 5) else WD_ALIGN_PARAGRAPH.JUSTIFY
            _cell_text(t.rows[i].cells[j], val, align=al)
    return t


def matrix_table(header, rows):
    """Tabel perbandingan n-kolom (Normal Table)."""
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        _cell_text(t.rows[0].cells[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            _cell_text(t.rows[i].cells[j], val)
    return t


# =====================================================================
#  ISI DOKUMEN
# =====================================================================
import content_bab  # noqa: E402
content_bab.build(globals())

doc.save(OUT)
print("Saved", OUT)
